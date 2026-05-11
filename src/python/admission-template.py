import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

def suffix(day):
    if 11 <= day <= 13:
        return "th"
    else:
        return {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")

date = datetime.now()

# Format the date as "21st March 2023"
formatted_date = date.strftime("%d{S} %B %Y").replace("{S}", suffix(date.day))


# Define the placeholders
placeholders = {
    "{{firstname}}": "JEMIMA",
    "{{middlename}}": "AFUA",
    "{{lastname}}": "KUMAH",
    "{{hometown}}": "DAMBAI",
    "{{region}}": "OTI",
    "{{datetoday}}": formatted_date,
    "{{programme}}": "Diploma in Jewellery Design Technology",
    "{{semesterstartmonth}}": "September",
    "{{semesterstartyear}}": "2023",
    "{{provisionalstudentnumber}}": "670",
    "{{firstpaymentdeadline}}": "17th July, 2023",
    "{{registrationperiod}}": "17th - 28th July, 2023",
    "{{orientationdate}}": "28th July 2023",
    "{{lecturesbegin}}": "Monday, 31st July 2023",
    "{{registrationstartdate}}": "31st July",
    "{{registrationenddate}}": "17th August 2023"
}

# Create a new document
document = docx.Document()

# Add the header
document.add_heading("", 0)

# create a new style based on an existing style
new_style = document.styles.add_style('New Style', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
new_style.base_style = document.styles['Normal']

# modify the properties of the new style
new_style.font.name = 'Times New Roman'
new_style.font.size = docx.shared.Pt(9)
new_style.font.bold = False

# Add the applicant's details
applicant_details = f"{placeholders['{{firstname}}']} {placeholders['{{middlename}}']} {placeholders['{{lastname}}']}\t\t\t\t\t\t {placeholders['{{datetoday}}']}\n{placeholders['{{hometown}}']}\n{placeholders['{{region}}']} "
paragraph = document.add_paragraph(applicant_details)
paragraph.style = new_style

# Add the offer details
offer_details = f"OFFER OF ADMISSION FOR THE {placeholders['{{programme}}']} PROGRAMME"
paragraph = document.add_paragraph()
run = paragraph.add_run(offer_details)
run.bold = True
run.font.all_caps = True

# Add the central details
central_details = f"The Central Admissions Committee has considered your application and is delighted to offer you admission at Level 100 of the "
programme = placeholders['{{programme}}']
rest_of_details = f" Programme commencing {placeholders['{{semesterstartmonth}}']} {placeholders['{{semesterstartyear}}']}. Your Provisional Student Number is "
provisional_student_number = placeholders['{{provisionalstudentnumber}}']
end_of_details = f".\n\nIf at any time the University discovers that you do not possess the qualifications by virtue of which you have been given admission into your programme of study, you will be withdrawn."
paragraph = document.add_paragraph()
for i, part in enumerate(central_details.split(programme)):
    run = paragraph.add_run(part)
    if i == 1:
        run.bold = True
run = paragraph.add_run(programme)
run.bold = True
run.font.all_caps = True
for i, part in enumerate(rest_of_details.split(provisional_student_number)):
    run = paragraph.add_run(part)
    if i == 1:
        run.bold = True
run = paragraph.add_run(provisional_student_number)
run.bold = True
run = paragraph.add_run(end_of_details)

# create a table with two rows and two columns
table = document.add_table(rows=5, cols=2)
table.columns[0].width = Pt(3)
table.columns[1].width = Pt(2)

row = table.rows[0]
row.cells[0].text = "\tTuition Fees:"
value_cell = row.cells[1]
value_cell_text = "GHȼ 1,234.50 per semester"

# Split the cell's text into two parts
value_text_parts = value_cell_text.split(' ')
value_cell.text = value_text_parts[0] + ' ' + value_text_parts[1]

# Add the rest of the text to the cell as a new run
value_cell.paragraphs[0].add_run(' ')

for i, part in enumerate(value_text_parts[2:]):
    # add non-breaking space before the word
    value_cell.paragraphs[0].add_run(u"\u00A0")
    # add the word
    value_cell.paragraphs[0].add_run(part)
    # apply bold to the first two runs only
    if i < 2:
        value_cell.paragraphs[0].runs[i].bold = True

row = table.rows[1]
row.cells[0].text = "\tRegistration and other Fees:"
value_cell = row.cells[1]
value_cell.text = "GHȼ 702.50"
value_cell.paragraphs[0].runs[0].bold = True

row = table.rows[2]
row.cells[0].text = "\tMinimum fees to be paid before \n\tRegistration (75%):"
value_cell = row.cells[1]
value_cell.text = "GHȼ 1,452.75"
value_cell.paragraphs[0].runs[0].bold = True

row = table.rows[3]
row.cells[0].text = "\tBalance to be paid before end-of-\n\tsemester exams (25%):"
value_cell = row.cells[1]
value_cell.text = "GHȼ 484.25"
value_cell.paragraphs[0].runs[0].bold = True

row = table.rows[4]
cell = row.cells[0]
cell.text = "\tHostel: The University has a decent and affordable hostel for students."
cell.merge(row.cells[1])

payment_details = f"If you accept this offer of admission, then you are required to pay not less than 75% of the total fees due at any Consolidated Bank Ghana, or the Kokomlemle Branch, into "
paragraph = document.add_paragraph(payment_details)
# Create the paragraph and add the runs
paragraph = document.add_paragraph()
paragraph.add_run("Account Number: 0600779100001").bold = True
paragraph.add_run(" by ")
paragraph.add_run(placeholders['{{firstpaymentdeadline}}']).bold = True
paragraph.add_run(". Full payment of fees for the Semester is expected before the End-of-Semester examinations begin. Please quote your Provisional Student Number when making payment. All fresh students are required to take a medical examination.")

# Add the first paragraph
document.add_paragraph("Arrangements for the semester are as follows:")

# Add the second paragraph with a placeholder for registration period
document.add_paragraph(f"\ti. Registration\t\t\t\t - \t{placeholders['{{registrationperiod}}']}")

# Add the third paragraph with a placeholder for orientation date
document.add_paragraph(f"\tii. Orientation for New Students\t - \t{placeholders['{{orientationdate}}']}")

# Add the fourth paragraph with a placeholder for the start of lectures
document.add_paragraph(f"\tiii. Lectures begin\t\t\t - \t{placeholders['{{lecturesbegin}}']}")

# Create the paragraph and add the runs
paragraph = document.add_paragraph()
paragraph.add_run("Students will be allowed to register between ").bold = False
paragraph.add_run(placeholders['{{registrationstartdate}}']).bold = True
paragraph.add_run(" and ").bold = False
paragraph.add_run(placeholders['{{registrationenddate}}']).bold = True
paragraph.add_run(" but only on payment of a late registration fee. No student will be allowed to register after ").bold = False
paragraph.add_run(placeholders['{{registrationenddate}}']).bold = True
paragraph.add_run(".")

document.add_paragraph("You will be on probation for the full duration of your programme and may be dismissed at any time for unsatisfactory academic work or grievous misconduct. You will be required to adhere to ALL University rules and regulations as contained in the Student Handbook, a copy of which will be made available to you during orientation.")

# Create the paragraph and add the runs
paragraph = document.add_paragraph()
paragraph.add_run("Please note that fees once paid are non-refundable. ").bold = False
paragraph.add_run("NO REFUND").bold = True
paragraph.add_run(" will be made in the event of a student's withdrawal or suspension from the University. Students are therefore advised to ensure that they complete their programmes of study within the prescribed period.")
document.add_paragraph("Once again, please accept our congratulations. We look forward to seeing you.")
document.add_paragraph(f"\n\n\nRoland A. Assale\nRegistrar")

filename = provisional_student_number + " AUCDT NEW ADMISSION LETTER.docx"
document.save(filename)

# Open the file using the default application
os.startfile(filename)
